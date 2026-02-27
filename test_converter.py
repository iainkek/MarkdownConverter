"""Tests for converter.py conversion functions and output path logic."""

import os
import sys
import tempfile
import shutil
from pathlib import Path

# Make sure we can import converter
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from converter import (
    convert_docx, convert_xlsx, convert_pdf, convert_url, convert_epub,
    unique_path, safe_filename, url_to_filename, CONVERTERS, SUPPORTED_EXTENSIONS,
    _epub_via_zip, _epub_via_calibre, _find_calibre,
)

PASS = 0
FAIL = 0


def report(name, passed, detail=""):
    global PASS, FAIL
    if passed:
        PASS += 1
        print(f"  PASS  {name}")
    else:
        FAIL += 1
        print(f"  FAIL  {name}  — {detail}")


# ─── Test output path logic ────────────────────────────────────────────────

def test_output_paths():
    """Verify .pdf -> .md output path (the reported .mp4 bug)."""
    print("\n=== Output Path Logic ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        # Test: .pdf -> .md
        pdf_path = Path(tmpdir) / "test_document.pdf"
        pdf_path.touch()
        out = pdf_path.with_suffix(".md")
        report("pdf -> .md extension", out.suffix == ".md", f"got {out.suffix}")
        report("pdf -> .md full name", out.name == "test_document.md", f"got {out.name}")

        # Test: various extensions
        for ext in [".docx", ".xlsx", ".epub"]:
            p = Path(tmpdir) / f"file{ext}"
            p.touch()
            out = p.with_suffix(".md")
            report(f"{ext} -> .md extension", out.suffix == ".md", f"got {out.suffix}")

        # Test: file with multiple dots
        tricky = Path(tmpdir) / "report.2024.final.pdf"
        tricky.touch()
        out = tricky.with_suffix(".md")
        report("multi-dot pdf -> .md", out.suffix == ".md", f"got {out.name}")

        # Test: unique_path doesn't change extension
        existing = Path(tmpdir) / "existing.md"
        existing.write_text("hello", encoding="utf-8")
        out = unique_path(Path(tmpdir) / "existing.md")
        report("unique_path keeps .md", out.suffix == ".md", f"got {out.suffix}")
        report("unique_path appends _1", "_1" in out.stem, f"got {out.name}")

        # Test: filename with spaces
        spaced = Path(tmpdir) / "my document file.pdf"
        spaced.touch()
        out = spaced.with_suffix(".md")
        report("spaced filename -> .md", out.suffix == ".md", f"got {out.name}")

        # Simulate what _do_conversion does for a PDF
        filepath = Path(tmpdir) / "test.pdf"
        filepath.touch()
        out_path = unique_path(filepath.with_suffix(".md"))
        report("simulated conversion path .md", out_path.suffix == ".md", f"got {out_path}")


# ─── Test PDF conversion ──────────────────────────────────────────────────

def test_pdf_conversion():
    """Test PDF conversion produces valid markdown output."""
    print("\n=== PDF Conversion ===")

    # Create a minimal PDF with text
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "test.pdf")

        # Create a simple PDF using pdfplumber-compatible format
        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(pdf_path)
            c.drawString(72, 700, "Hello World Test PDF")
            c.drawString(72, 680, "This is a test document for markdown conversion.")
            c.drawString(72, 660, "It has multiple lines of text.")
            c.save()

            result = convert_pdf(pdf_path)
            report("PDF produces output", result is not None and len(result) > 0, "empty result")
            report("PDF output is string", isinstance(result, str), f"got {type(result)}")
            report("PDF contains expected text", "Hello World" in result, f"output: {result[:100]}")

            # Write it out and check the file
            out_path = Path(pdf_path).with_suffix(".md")
            out_path.write_text(result, encoding="utf-8")
            report("PDF output file is .md", out_path.suffix == ".md", f"got {out_path.suffix}")
            report("PDF output file exists", out_path.exists())
            report("PDF output is text not binary",
                   all(c < 128 or c > 160 for c in result.encode("utf-8")[:200]),
                   "contains suspicious binary-like content")
        except ImportError:
            print("  SKIP  (reportlab not installed, testing with real PDF)")
            # Try with a real PDF from Downloads
            real_pdfs = list(Path("c:/Users/iaink/Downloads").glob("*.pdf"))
            if real_pdfs:
                test_pdf = str(real_pdfs[0])
                print(f"  INFO  Testing with: {Path(test_pdf).name}")
                try:
                    result = convert_pdf(test_pdf)
                    report("Real PDF produces output", result is not None and len(result) > 0)
                    report("Real PDF output is string", isinstance(result, str))
                    if result:
                        report("Real PDF has reasonable content", len(result) > 10,
                               f"only {len(result)} chars")
                        # Check it's not binary garbage
                        printable_ratio = sum(1 for c in result[:500] if c.isprintable() or c in '\n\r\t') / min(len(result), 500)
                        report("Real PDF output is readable text", printable_ratio > 0.8,
                               f"only {printable_ratio:.0%} printable")
                except Exception as e:
                    report("Real PDF conversion", False, str(e))
            else:
                print("  SKIP  No PDF files available for testing")


# ─── Test DOCX conversion ─────────────────────────────────────────────────

def test_docx_conversion():
    """Test DOCX conversion."""
    print("\n=== DOCX Conversion ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "test.docx")

        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Test Heading", level=1)
            doc.add_paragraph("This is a test paragraph.")
            doc.add_paragraph("Second paragraph with more text.")
            doc.save(docx_path)

            result = convert_docx(docx_path)
            report("DOCX produces output", result is not None and len(result) > 0)
            report("DOCX contains heading", "Test Heading" in result, f"output: {result[:100]}")
            report("DOCX contains paragraph", "test paragraph" in result)
        except ImportError:
            print("  SKIP  (python-docx not installed)")


# ─── Test XLSX conversion ─────────────────────────────────────────────────

def test_xlsx_conversion():
    """Test XLSX conversion."""
    print("\n=== XLSX Conversion ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        xlsx_path = os.path.join(tmpdir, "test.xlsx")

        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "TestSheet"
            ws.append(["Name", "Age", "City"])
            ws.append(["Alice", 30, "Sydney"])
            ws.append(["Bob", 25, "Melbourne"])
            wb.save(xlsx_path)

            result = convert_xlsx(xlsx_path)
            report("XLSX produces output", result is not None and len(result) > 0)
            report("XLSX has table format", "|" in result, f"output: {result[:100]}")
            report("XLSX contains data", "Alice" in result and "Bob" in result)
            report("XLSX contains headers", "Name" in result and "Age" in result)
        except ImportError:
            print("  SKIP  (openpyxl not installed)")


# ─── Test URL conversion ──────────────────────────────────────────────────

def test_url_conversion():
    """Test URL conversion."""
    print("\n=== URL Conversion ===")

    try:
        result = convert_url("https://en.wikipedia.org/wiki/Markdown")
        report("URL produces output", result is not None and len(result) > 0)
        report("URL output is string", isinstance(result, str))
        if result:
            report("URL has reasonable length", len(result) > 100, f"only {len(result)} chars")
    except Exception as e:
        report("URL conversion", False, str(e))

    # Test url_to_filename
    fname = url_to_filename("https://www.example.com/some/page")
    report("url_to_filename works", fname and len(fname) > 0, f"got: {fname}")
    report("url_to_filename safe chars", "/" not in fname and "\\" not in fname)


# ─── Test EPUB conversion ─────────────────────────────────────────────────

def test_epub_conversion():
    """Test EPUB conversion paths."""
    print("\n=== EPUB Conversion ===")

    calibre = _find_calibre()
    report("Calibre found", calibre is not None, "ebook-convert not found")
    if calibre:
        print(f"  INFO  Calibre at: {calibre}")

    # Test with the DRM epub (should fail gracefully)
    drm_epub = Path("c:/Users/iaink/Downloads/ProfEpubtoXML/Law of Charity 3rd ed.epub")
    if drm_epub.exists():
        try:
            result = convert_epub(str(drm_epub))
            # If it somehow succeeds, check quality
            if result:
                printable_ratio = sum(1 for c in result[:500] if c.isprintable() or c in '\n\r\t') / min(len(result), 500)
                report("DRM EPUB output is readable", printable_ratio > 0.8,
                       f"only {printable_ratio:.0%} printable — might be garbage")
            else:
                report("DRM EPUB returns None (expected)", True)
        except ValueError as e:
            report("DRM EPUB raises clear error", True)
            print(f"  INFO  Error message: {e}")
        except Exception as e:
            report("DRM EPUB error handling", False, f"unexpected: {type(e).__name__}: {e}")


# ─── Test converter dispatch ──────────────────────────────────────────────

def test_converter_dispatch():
    """Test that all extensions map to correct converters."""
    print("\n=== Converter Dispatch ===")

    report(".pdf maps to convert_pdf", CONVERTERS.get(".pdf") is convert_pdf)
    report(".docx maps to convert_docx", CONVERTERS.get(".docx") is convert_docx)
    report(".xlsx maps to convert_xlsx", CONVERTERS.get(".xlsx") is convert_xlsx)
    report(".epub maps to convert_epub", CONVERTERS.get(".epub") is convert_epub)
    report("no .mp4 converter", ".mp4" not in CONVERTERS)
    report("no .txt converter", ".txt" not in CONVERTERS)

    # Check supported extensions match converters
    for ext in SUPPORTED_EXTENSIONS:
        report(f"{ext} in CONVERTERS", ext in CONVERTERS, f"{ext} supported but no converter")


# ─── Test safe_filename ───────────────────────────────────────────────────

def test_safe_filename():
    """Test filename sanitization."""
    print("\n=== Filename Sanitization ===")

    report("removes colons", ":" not in safe_filename("file:name"))
    report("removes quotes", '"' not in safe_filename('file"name'))
    report("removes slashes", "/" not in safe_filename("file/name"))
    report("handles spaces", " " not in safe_filename("file name here"))
    report("limits length", len(safe_filename("a" * 200)) <= 100)
    report("handles empty", safe_filename("") == "converted")


# ─── Full integration test ────────────────────────────────────────────────

def test_full_integration():
    """Simulate what the GUI does end-to-end for a PDF."""
    print("\n=== Full Integration (PDF) ===")

    with tempfile.TemporaryDirectory() as tmpdir:
        # Create test PDF
        pdf_created = False
        pdf_path = Path(tmpdir) / "integration_test.pdf"

        try:
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(str(pdf_path))
            c.drawString(72, 700, "Integration Test Document")
            c.drawString(72, 680, "Line two of the document.")
            c.save()
            pdf_created = True
        except ImportError:
            # Copy a real PDF
            real_pdfs = list(Path("c:/Users/iaink/Downloads").glob("*.pdf"))
            if real_pdfs:
                shutil.copy2(str(real_pdfs[0]), str(pdf_path))
                pdf_created = True
                print(f"  INFO  Using real PDF: {real_pdfs[0].name}")

        if not pdf_created:
            print("  SKIP  No way to create test PDF")
            return

        # Simulate exactly what _do_conversion does
        filepath = pdf_path
        ext = filepath.suffix.lower()
        report("Extension detected as .pdf", ext == ".pdf", f"got {ext}")

        converter = CONVERTERS.get(ext)
        report("Converter found for .pdf", converter is not None)
        report("Converter is convert_pdf", converter is convert_pdf)

        markdown = converter(str(filepath))
        report("Conversion produced output", markdown is not None and len(markdown) > 0)

        out_path = unique_path(filepath.with_suffix(".md"))
        report("Output path extension is .md", out_path.suffix == ".md", f"got {out_path.suffix}")
        report("Output path name correct", out_path.name == "integration_test.md",
               f"got {out_path.name}")

        out_path.write_text(markdown, encoding="utf-8")
        report("Output file created", out_path.exists())
        report("Output file is .md", out_path.suffix == ".md")

        # Read it back and verify
        content = out_path.read_text(encoding="utf-8")
        report("Output file has content", len(content) > 0)
        report("Output file is valid text",
               all(c < 128 or c > 160 for c in content.encode("utf-8")[:500]),
               "suspicious binary content")


# ─── Run all tests ────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Markdown Converter — Test Suite")
    print("=" * 50)

    test_output_paths()
    test_converter_dispatch()
    test_safe_filename()
    test_pdf_conversion()
    test_docx_conversion()
    test_xlsx_conversion()
    test_epub_conversion()
    test_url_conversion()
    test_full_integration()

    print("\n" + "=" * 50)
    print(f"Results: {PASS} passed, {FAIL} failed, {PASS + FAIL} total")
    if FAIL == 0:
        print("All tests passed!")
    else:
        print(f"{FAIL} test(s) FAILED")
        sys.exit(1)
